from fastapi import FastAPI, UploadFile, File, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from fastapi.responses import FileResponse
from typing import Optional
from collections import Counter
import requests
from main_controller import *
import shutil
import nltk
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import os
from os.path import exists


nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')

app = FastAPI()

origins = [
    "http://localhost:3000",
    "https://localhost:3000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
class TranslationSchema(BaseModel):
    text:str
 

from googletrans import Translator
translator = Translator()


def translate_text(inputText): 
    detectLang= translator.detect(inputText).lang
    return translator.translate(inputText, src=detectLang, dest="hi").text
 
def audio_to_text(filename):
    print("Filename", filename)
    audio_url = upload(filename)
    print("Audio upoad done")
    save_transcript(audio_url, 'file_title')
    print("transcription done")

def create_word():
    text_file = open("file_title.txt", "r")
    data = text_file.read()
    text_file.close()
    document = Document()
    document.add_picture('logo.png', width=Inches(1.25))

    document.add_heading('Notes', 0)

    p = document.add_paragraph(data)

    document.add_page_break()

    document.save('demo.docx')


@app.get("/")
def root():
    return{"greeting":"Hey."}

@app.post("/translation")
def translation(req: TranslationSchema):
    # text = "This will be fetchced by api."
    translated_text = translate_text(req.text)
    return{"text" : translated_text}

@app.get("/word")
def word():
    file_exists = exists("file_title.txt")
    if file_exists:
        create_word()
        return FileResponse("demo.docx", media_type="application/msword", filename="Lecture.docx")
    else:
        return{"status":"failed"}

@app.get("/txt")
def txt():
    file_exists = exists("file_title.txt")
    if file_exists:
        return FileResponse("file_title.txt", media_type="text/plain", filename="Lecture.txt")
    else:
        return{"status":"failed"}


@app.get("/pdf")
def pdf():
    file_exists = exists("file_title.txt")
    if file_exists:
        create_word()
        convert("demo.docx")
        return FileResponse("demo.pdf", media_type="application/pdf", filename="Lecture.pdf")
    else:
        return{"status":"failed"}




    
@app.post("/upload")
async def uploadFile(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    with open(f'{file.filename}', "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    background_tasks.add_task(audio_to_text, file.filename)
    return{"name":file.filename}



def classify(file):    
    text_file = open(file, "r")
    data = text_file.read()
    text_file.close()

    tokens = nltk.word_tokenize(data.lower())
    data = nltk.Text(tokens)
    tags = nltk.pos_tag(data)
    dataset = Counter(tag for word,tag in tags)
    return dataset
    
@app.get("/polling")
def polling():
    file_exists = exists("file_title.txt")
    if file_exists:
        dataset = classify("file_title.txt")
    return{"data":dataset, "resolved":True}
    

    
    
  

    

